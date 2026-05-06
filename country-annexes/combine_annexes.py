#!/usr/bin/env python3
"""
Combine four scenario DOCX files into one per fund size.

Creates a single Word document per fund size containing all four scenarios:
  iusaf-pure, iusaf-strict, gini-minimum, band-order-boundary.

Usage:
    python3 country-annexes/combine_annexes.py
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

FONT = "Times New Roman"
HEADER_BG = "D9D9D9"
SIDS_HIGHLIGHT = "E8F5E9"

SCENARIOS = ['iusaf-pure', 'iusaf-strict', 'gini-minimum', 'band-order-boundary']
SCENARIO_TITLES = {
    'iusaf-pure': ('IUSAF (Pure)', 'Pure IUSAF allocation with band inversion, no TSAC or SOSAC'),
    'iusaf-strict': ('Strict (TSAC 1.5%, SOSAC 3%)', 'Strict balance point'),
    'gini-minimum': ('Gini-minimum (TSAC 2.5%, SOSAC 3%)', 'Gini-minimum balance point'),
    'band-order-boundary': ('Band-order boundary (TSAC 3.0%, SOSAC 3%)', 'Band-order boundary'),
}

DISPLAY_COLS = [
    'party', 'total_allocation', 'state_component', 'iplc_component',
    'component_iusaf_amt', 'component_tsac_amt', 'component_sosac_amt',
    'WB Income Group', 'is_ldc', 'is_sids', 'is_eu_ms', 'un_band'
]

FUND_DIRS = [
    ('fifty-million', '50M'),
    ('two-hundred-million', '200M'),
    ('five-hundred-million', '500M'),
    ('one-billion', '1B'),
]


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _gini(values):
    v = np.sort(values)
    n = len(v)
    if n == 0 or np.sum(v) == 0:
        return 0.0
    idx = np.arange(1, n + 1)
    return float((2 * np.sum(idx * v) - (n + 1) * np.sum(v)) / (n * np.sum(v)))


def build_combined_docx(fund_label, fund_display, fund_dir):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Title
    title = doc.add_heading(f'Country Annexes: USD {fund_display} Fund', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = FONT
        run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f'142 eligible Parties | State/IPLC: 50/50 | '
        f'Four balance-point scenarios ranked by total allocation')
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.italic = True

    for sid in SCENARIOS:
        sname, sdesc = SCENARIO_TITLES[sid]
        csv_path = os.path.join(fund_dir, sid, f'{sid}-country-annex.csv')
        if not os.path.exists(csv_path):
            print(f'  WARNING: {csv_path} not found, skipping')
            continue

        df = pd.read_csv(csv_path)
        doc.add_page_break()

        # Section heading
        h = doc.add_heading(f'{sname}', level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.name = FONT
            run.font.size = Pt(13)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(sdesc)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.italic = True

        # Summary stats
        total_alloc = df['total_allocation'].sum()
        total_iusaf = df['component_iusaf_amt'].sum()
        total_tsac = df['component_tsac_amt'].sum()
        total_sosac = df['component_sosac_amt'].sum()
        ldc_total = df[df['is_ldc'] == True]['total_allocation'].sum()
        sids_total = df[df['is_sids'] == True]['total_allocation'].sum()
        n_eligible = len(df)
        gini = _gini(df['total_allocation'].values)

        p = doc.add_paragraph()
        run = p.add_run(
            f'{n_eligible} eligible Parties | Fund: USD {fund_display} | '
            f'State/IPLC: 50/50 | Gini: {gini:.4f} | '
            f'IUSAF: {total_iusaf:.1f} M ({total_iusaf/total_alloc*100:.1f}%), '
            f'TSAC: {total_tsac:.1f} M ({total_tsac/total_alloc*100:.1f}%), '
            f'SOSAC: {total_sosac:.1f} M ({total_sosac/total_alloc*100:.1f}%), '
            f'LDC: {ldc_total:.1f} M, SIDS: {sids_total:.1f} M.'
        )
        run.font.name = FONT
        run.font.size = Pt(8.5)

        doc.add_paragraph()

        # Table
        word_headers = [
            'Rank', 'Party', 'Total\n(M)', 'State\n(M)', 'IPLC\n(M)',
            'IUSAF\n(M)', 'TSAC\n(M)', 'SOSAC\n(M)',
            'Income', 'LDC', 'SIDS', 'EU', 'Band'
        ]
        n_cols = len(word_headers)
        n_rows = len(df) + 2  # header + data + total
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header
        for j, hdr in enumerate(word_headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p_h = cell.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for k, line in enumerate(hdr.split("\n")):
                if k > 0:
                    p_h.add_run("\n")
                run = p_h.add_run(line)
                run.font.name = FONT
                run.font.size = Pt(6.5)
                run.font.bold = True
            set_cell_shading(cell, HEADER_BG)

        # Data rows
        for i, (idx, row) in enumerate(df.iterrows()):
            is_sids = row.get('is_sids', False)
            values = [
                str(i + 1),
                str(row['party']),
                f"{row['total_allocation']:.2f}",
                f"{row['state_component']:.2f}",
                f"{row['iplc_component']:.2f}",
                f"{row['component_iusaf_amt']:.2f}",
                f"{row['component_tsac_amt']:.2f}",
                f"{row['component_sosac_amt']:.2f}",
                str(row['WB Income Group']),
                'LDC' if row['is_ldc'] else '\u2013',
                'SIDS' if row['is_sids'] else '\u2013',
                'EU' if row['is_eu_ms'] else '\u2013',
                str(row['un_band']).split(':')[0] if pd.notna(row['un_band']) else '',
            ]
            for j, val in enumerate(values):
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                p_d = cell.paragraphs[0]
                p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 2, 3, 4, 5, 6, 7] else WD_ALIGN_PARAGRAPH.LEFT
                run = p_d.add_run(val)
                run.font.name = FONT
                run.font.size = Pt(6.5)
                if j == 1 and is_sids:
                    run.font.bold = True
            if is_sids:
                for j in range(n_cols):
                    set_cell_shading(table.rows[i + 1].cells[j], SIDS_HIGHLIGHT)

        # Total row
        total_vals = [
            '', 'Total',
            f"{total_alloc:.2f}", f"{df['state_component'].sum():.2f}", f"{df['iplc_component'].sum():.2f}",
            f"{total_iusaf:.2f}", f"{total_tsac:.2f}", f"{total_sosac:.2f}",
            '', '', '', '', '',
        ]
        for j, val in enumerate(total_vals):
            cell = table.rows[len(df) + 1].cells[j]
            cell.text = ""
            p_t = cell.paragraphs[0]
            p_t.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in [2, 3, 4, 5, 6, 7] else WD_ALIGN_PARAGRAPH.LEFT
            run = p_t.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(6.5)
            run.font.bold = True
            set_cell_shading(cell, HEADER_BG)

    out_path = os.path.join(fund_dir, f'country-annexes-{fund_label}.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')


def main():
    base = os.path.dirname(os.path.abspath(__file__))
    for fund_label, fund_display in FUND_DIRS:
        fund_dir = os.path.join(base, fund_label)
        print(f'\nCombining: {fund_label} ({fund_display})')
        build_combined_docx(fund_label, fund_display, fund_dir)

    print('\nDone.')


if __name__ == "__main__":
    main()
