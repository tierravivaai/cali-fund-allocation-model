#!/usr/bin/env python3
"""
Generate Country Annex Tables
==============================

Produces per-country allocation tables for four IUSAF scenarios,
drawn directly from the application calculator to ensure consistency
with the live model.

Four scenarios:
  a) IUSAF (pure):     beta=0, gamma=0 — baseline band-inversion allocation
  b) Strict:            beta=0.015, gamma=0.03 — TSAC 1.5%, SOSAC 3%
  c) Gini-minimum:      beta=0.025, gamma=0.03 — TSAC 2.5%, SOSAC 3%
  d) Band-order boundary: beta=0.03, gamma=0.03 — TSAC 3.0%, SOSAC 3%

Each scenario folder contains:
  - <scenario>-country-annex.csv   Machine-readable data
  - <scenario>-country-annex.md    Human-readable documentation
  - <scenario>-country-annex.docx Word-formatted table for the paper

Usage:
    python3 country-annexes/generate_country_annexes.py
"""

import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

import duckdb
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from cali_model.data_loader import load_data, get_base_data
from cali_model.calculator import calculate_allocations

# ── Configuration ────────────────────────────────────────────────────────────

FUND = 1_000_000_000
IPLC = 50
EXCLUDE_HI = True
HI_MODE = "exclude_except_sids"
UN_SCALE = "band_inversion"

FONT = "Times New Roman"
HEADER_BG = "D9D9D9"
SIDS_HIGHLIGHT = "E8F5E9"

SCENARIOS = [
    {
        "id": "iusaf-pure",
        "name": "IUSAF (Pure)",
        "short": "Pure IUSAF",
        "beta": 0.0,
        "gamma": 0.0,
        "description": "Pure IUSAF allocation with band inversion, no TSAC or SOSAC",
    },
    {
        "id": "iusaf-strict",
        "name": "Strict",
        "short": "Strict (1.5/3)",
        "beta": 0.015,
        "gamma": 0.03,
        "description": "Strict balance point: TSAC 1.5%, SOSAC 3%",
    },
    {
        "id": "gini-minimum",
        "name": "Gini-minimum",
        "short": "Gini-min (2.5/3)",
        "beta": 0.025,
        "gamma": 0.03,
        "description": "Gini-minimum balance point: TSAC 2.5%, SOSAC 3%",
    },
    {
        "id": "band-order-boundary",
        "name": "Band-order boundary",
        "short": "Boundary (3.0/3)",
        "beta": 0.03,
        "gamma": 0.03,
        "description": "Band-order boundary: TSAC 3.0%, SOSAC 3%",
    },
]

DISPLAY_COLS = [
    'party', 'total_allocation', 'state_component', 'iplc_component',
    'component_iusaf_amt', 'component_tsac_amt', 'component_sosac_amt',
    'WB Income Group', 'is_ldc', 'is_sids', 'is_eu_ms', 'un_band'
]
CSV_HEADERS = [
    'Rank', 'Party', 'Total (USD M)', 'State (USD M)', 'IPLC (USD M)',
    'IUSAF (USD M)', 'TSAC (USD M)', 'SOSAC (USD M)',
    'Income Group', 'LDC', 'SIDS', 'EU MS', 'UN Band'
]


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def generate_scenario(con, scenario):
    """Generate CSV, MD, and DOCX for one scenario."""
    sid = scenario["id"]
    sname = scenario["name"]
    beta = scenario["beta"]
    gamma = scenario["gamma"]
    
    base_df = get_base_data(con)
    df = calculate_allocations(
        base_df, FUND, IPLC,
        exclude_high_income=EXCLUDE_HI,
        high_income_mode=HI_MODE,
        tsac_beta=beta, sosac_gamma=gamma,
        equality_mode=False,
        un_scale_mode=UN_SCALE,
    )
    
    # Filter to eligible, sort by allocation desc then party name asc
    eligible = df[df['eligible']].copy()
    eligible = eligible.sort_values(
        ['total_allocation', 'party'], ascending=[False, True]
    ).reset_index(drop=True)
    eligible.index = eligible.index + 1
    eligible.index.name = 'Rank'
    
    out_dir = os.path.join(os.path.dirname(__file__), sid)
    os.makedirs(out_dir, exist_ok=True)
    
    # ── CSV ──────────────────────────────────────────────────────────────
    csv_df = eligible[DISPLAY_COLS].copy()
    csv_df.insert(0, 'Rank', csv_df.index)
    csv_path = os.path.join(out_dir, f"{sid}-country-annex.csv")
    csv_df.to_csv(csv_path, index=False, float_format='%.4f')
    print(f"  Saved: {csv_path} ({len(csv_df)} rows)")
    
    # ── Calculate summary stats ──────────────────────────────────────────
    n_eligible = len(eligible)
    total_alloc = eligible['total_allocation'].sum()
    total_state = eligible['state_component'].sum()
    total_iplc = eligible['iplc_component'].sum()
    total_iusaf = eligible['component_iusaf_amt'].sum()
    total_tsac = eligible['component_tsac_amt'].sum()
    total_sosac = eligible['component_sosac_amt'].sum()
    ldc_total = eligible[eligible['is_ldc']]['total_allocation'].sum()
    sids_total = eligible[eligible['is_sids']]['total_allocation'].sum()
    gini = _gini(eligible['total_allocation'].values)
    
    # ── Markdown ─────────────────────────────────────────────────────────
    md_path = os.path.join(out_dir, f"{sid}-country-annex.md")
    with open(md_path, 'w') as f:
        f.write(f"# Country Annex: {sname}\n\n")
        f.write(f"**{scenario['description']}**\n\n")
        f.write(f"| Parameter | Value |\n|-----------|-------|\n")
        f.write(f"| Fund size | USD {FUND/1e6:.0f} million |\n")
        f.write(f"| State/IPLC split | {IPLC}/{100-IPLC} |\n")
        f.write(f"| TSAC (beta) | {beta*100:.1f}% |\n")
        f.write(f"| SOSAC (gamma) | {gamma*100:.1f}% |\n")
        f.write(f"| Eligible parties | {n_eligible} |\n")
        f.write(f"| Gini coefficient | {gini:.4f} |\n")
        f.write(f"| LDC total | USD {ldc_total:.2f} M |\n")
        f.write(f"| SIDS total | USD {sids_total:.2f} M |\n\n")
        f.write(f"## Totals\n\n")
        f.write(f"| Component | USD M | Share |\n")
        f.write(f"|-----------|-------|-------|\n")
        f.write(f"| IUSAF | {total_iusaf:.2f} | {total_iusaf/total_alloc*100:.1f}% |\n")
        f.write(f"| TSAC | {total_tsac:.2f} | {total_tsac/total_alloc*100:.1f}% |\n")
        f.write(f"| SOSAC | {total_sosac:.2f} | {total_sosac/total_alloc*100:.1f}% |\n")
        f.write(f"| **Total** | **{total_alloc:.2f}** | **100.0%** |\n\n")
        f.write(f"## Per-Country Allocations\n\n")
        f.write("See the accompanying CSV file for full data.\n\n")
        f.write(f"Generated by `generate_country_annexes.py` using the live calculator.\n")
    print(f"  Saved: {md_path}")
    
    # ── Word document ────────────────────────────────────────────────────
    doc_path = os.path.join(out_dir, f"{sid}-country-annex.docx")
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Title
    title = doc.add_heading(f"Country Annex: {sname}", level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = FONT
        run.font.size = Pt(12)
    
    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_text = (f"{n_eligible} eligible Parties | Fund: USD {FUND/1e6:.0f} M | "
                f"State/IPLC: {IPLC}/{100-IPLC} | Gini: {gini:.4f}")
    run = sub.add_run(sub_text)
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.italic = True
    
    # Component summary paragraph
    p = doc.add_paragraph()
    run = p.add_run(
        f"IUSAF: USD {total_iusaf:.1f} M ({total_iusaf/total_alloc*100:.1f}%), "
        f"TSAC: USD {total_tsac:.1f} M ({total_tsac/total_alloc*100:.1f}%), "
        f"SOSAC: USD {total_sosac:.1f} M ({total_sosac/total_alloc*100:.1f}%), "
        f"LDC total: USD {ldc_total:.1f} M, SIDS total: USD {sids_total:.1f} M."
    )
    run.font.name = FONT
    run.font.size = Pt(8.5)
    
    doc.add_paragraph()
    
    # Table
    word_headers = [
        'Rank', 'Party', 'Total\n(M)', 'State\n(M)', 'IPLC\n(M)',
        'IUSAF\n(M)', 'TSAC\n(M)', 'SOSAC\n(M)',
        'Income', 'LDC', 'SIDS', 'EU', 'Band'
    ]
    n_cols = len(word_headers)
    n_rows = len(eligible) + 1  # +1 for total row
    
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)  # +1 for header
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    
    # Header row
    for j, hdr in enumerate(word_headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for k, line in enumerate(hdr.split("\n")):
            if k > 0:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.name = FONT
            run.font.size = Pt(6.5)
            run.font.bold = True
        set_cell_shading(cell, HEADER_BG)
    
    # Data rows
    for i, (idx, row) in enumerate(eligible.iterrows()):
        is_sids = row.get('is_sids', False)
        values = [
            str(idx),
            str(row['party']),
            f"{row['total_allocation']:.2f}",
            f"{row['state_component']:.2f}",
            f"{row['iplc_component']:.2f}",
            f"{row['component_iusaf_amt']:.2f}",
            f"{row['component_tsac_amt']:.2f}",
            f"{row['component_sosac_amt']:.2f}",
            str(row['WB Income Group']),
            'LDC' if row['is_ldc'] else '-',
            'SIDS' if row['is_sids'] else '-',
            'EU' if row['is_eu_ms'] else '-',
            str(row['un_band']).split(':')[0] if pd.notna(row['un_band']) else '',
        ]
        
        for j, val in enumerate(values):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 2, 3, 4, 5, 6, 7] else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            run = p.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(6.5)
            if j == 1 and is_sids:
                run.font.bold = True
        
        if is_sids:
            for j in range(n_cols):
                set_cell_shading(table.rows[i + 1].cells[j], SIDS_HIGHLIGHT)
    
    # Total row
    total_row_idx = len(eligible) + 1
    total_vals = [
        '', 'Total',
        f"{total_alloc:.2f}", f"{total_state:.2f}", f"{total_iplc:.2f}",
        f"{total_iusaf:.2f}", f"{total_tsac:.2f}", f"{total_sosac:.2f}",
        '', '', '', '', '',
    ]
    for j, val in enumerate(total_vals):
        cell = table.rows[total_row_idx].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in [2, 3, 4, 5, 6, 7] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.name = FONT
        run.font.size = Pt(6.5)
        run.font.bold = True
        set_cell_shading(cell, HEADER_BG)
    
    doc.save(doc_path)
    print(f"  Saved: {doc_path}")
    
    return eligible


def _gini(values):
    """Compute Gini coefficient."""
    import numpy as np
    v = np.sort(values)
    n = len(v)
    if n == 0 or np.sum(v) == 0:
        return 0.0
    idx = np.arange(1, n + 1)
    return float((2 * np.sum(idx * v) - (n + 1) * np.sum(v)) / (n * np.sum(v)))


def main():
    print("Generating country annex tables...")
    con = duckdb.connect(database=":memory:")
    load_data(con)
    
    for scenario in SCENARIOS:
        print(f"\n  {scenario['name']} (beta={scenario['beta']}, gamma={scenario['gamma']})")
        generate_scenario(con, scenario)
    
    print("\nDone. All four country annexes generated.")


if __name__ == "__main__":
    main()
