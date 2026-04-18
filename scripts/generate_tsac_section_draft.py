"""Generate the restructured TSAC sensitivity section as a Word document.

This produces a draft reference document for the paper revision, incorporating
Option D (band-order preservation replacing the Spearman 0.85 threshold).
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from cali_model.data_loader import load_data, get_base_data
from cali_model.calculator import calculate_allocations

FONT = "Times New Roman"
HEADER_BG = "D9D9D9"
OUT = os.path.join(os.path.dirname(__file__), '..', 'model-tables', 'iusaf-tsac-section-draft.docx')

FUND = 1_000_000_000
IPLC = 50


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    """Create a styled table with headers and data rows."""
    highlight_rows = highlight_rows or {}
    n_cols = len(headers)
    n_data = len(rows)
    table = doc.add_table(rows=n_data + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for k, line in enumerate(hdr.split("\n")):
            if k > 0:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.name = FONT
            run.font.size = Pt(8.5)
            run.font.bold = True
        set_cell_shading(cell, HEADER_BG)

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < len(row_data) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(8.5)

            if i in highlight_rows:
                color = highlight_rows[i]
                if color == "green":
                    run.font.bold = True
                    set_cell_shading(cell, "D4EDDA")
                elif color == "red":
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    run.font.bold = True

    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w

    return table


def compute_all_scenarios(con):
    """Compute metrics for all TSAC levels needed."""
    base_df = get_base_data(con)
    pure = calculate_allocations(base_df, FUND, IPLC, exclude_high_income=True,
                                 tsac_beta=0, sosac_gamma=0, equality_mode=False,
                                 un_scale_mode="band_inversion")
    pure_eligible = pure[pure['eligible']]

    results = {}
    tsac_levels = [0, 0.015, 0.025, 0.03, 0.035, 0.092]
    labels = {
        0: "Pure IUSAF",
        0.015: "Strict",
        0.025: "Gini-minimum",
        0.03: "Band-order boundary",
        0.035: "Bounded",
        0.092: "TSAC component overturn",
    }

    for beta in tsac_levels:
        gamma = 0.03
        df = calculate_allocations(base_df, FUND, IPLC, exclude_high_income=True,
                                   tsac_beta=beta, sosac_gamma=gamma, equality_mode=False,
                                   un_scale_mode="band_inversion")
        eligible = df[df['eligible']]

        merged = eligible[['party', 'final_share']].merge(
            pure_eligible[['party', 'final_share']], on='party', suffixes=('_cur', '_base'))
        r_cur = merged['final_share_cur'].rank(method='average')
        r_base = merged['final_share_base'].rank(method='average')
        spearman = float(r_cur.corr(r_base, method='pearson'))

        allocs = eligible['total_allocation'].values
        n = len(allocs)
        sorted_a = np.sort(allocs)
        gini = 2 * np.sum(np.arange(1, n+1) * sorted_a) / (n * sorted_a.sum()) - (n+1)/n

        b6 = eligible[eligible['un_band'].str.startswith('Band 6')]['total_allocation'].mean()
        b5 = eligible[eligible['un_band'].str.startswith('Band 5')]['total_allocation'].mean()
        margin = (b5 - b6) / b5 * 100 if b5 > 0 else 0
        order_ok = b5 > b6

        results[beta] = {
            'label': labels.get(beta, f"{beta*100:.1f}%"),
            'iusaf_pct': (1 - beta - gamma) * 100,
            'spearman': spearman,
            'gini': gini,
            'b6_mean': b6, 'b5_mean': b5,
            'margin': margin, 'order_ok': order_ok,
        }

    return results


def compute_trajectory(con):
    """Compute the full TSAC ranking trajectory."""
    base_df = get_base_data(con)
    pure = calculate_allocations(base_df, FUND, IPLC, exclude_high_income=True,
                                 tsac_beta=0, sosac_gamma=0, equality_mode=False,
                                 un_scale_mode="band_inversion")
    pure_eligible = pure[pure['eligible']]

    rows = []
    for beta_pct in [0, 0.5, 1.0, 1.5, 2.0, 2.5, 3.0, 3.5, 4.0, 4.5, 5.0, 6.0, 7.0, 8.0, 9.0, 9.2, 10.0]:
        beta = beta_pct / 100
        gamma = 0.03
        df = calculate_allocations(base_df, FUND, IPLC, exclude_high_income=True,
                                   tsac_beta=beta, sosac_gamma=gamma, equality_mode=False,
                                   un_scale_mode="band_inversion")
        eligible = df[df['eligible']]

        merged = eligible[['party', 'final_share']].merge(
            pure_eligible[['party', 'final_share']], on='party', suffixes=('_cur', '_base'))
        r_cur = merged['final_share_cur'].rank(method='average')
        r_base = merged['final_share_base'].rank(method='average')
        spearman = float(r_cur.corr(r_base, method='pearson'))

        allocs = eligible['total_allocation'].values
        n = len(allocs)
        sorted_a = np.sort(allocs)
        gini = 2 * np.sum(np.arange(1, n+1) * sorted_a) / (n * sorted_a.sum()) - (n+1)/n

        b6 = eligible[eligible['un_band'].str.startswith('Band 6')]['total_allocation'].mean()
        b5 = eligible[eligible['un_band'].str.startswith('Band 5')]['total_allocation'].mean()

        rows.append({
            'TSAC %': beta_pct,
            'IUSAF %': (1 - beta - gamma) * 100,
            'Spearman ρ': round(spearman, 3),
            'Gini': round(gini, 4),
            'Band 6 mean': round(b6, 2),
            'Band 5 mean': round(b5, 2),
            'Band order': 'Preserved' if b5 > b6 else 'Overturned',
        })
    return pd.DataFrame(rows)


def main():
    con = duckdb.connect(database=':memory:')
    load_data(con)
    scenarios = compute_all_scenarios(con)
    trajectory = compute_trajectory(con)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- HEADING ----
    h = doc.add_heading("TSAC Sensitivity and Balance Points (Option D Revision Draft)", level=1)
    for run in h.runs:
        run.font.name = FONT
        run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("DRAFT — For reference during paper revision\n142 eligible Parties | Fund size: USD 1,000 M | SOSAC = 3% | State/IPLC split: 50/50")
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ---- SECTION 1: Introduction ----
    h2 = doc.add_heading("Sensitivity Analysis: Band-Order Preservation as the Structural Constraint", level=2)
    for run in h2.runs:
        run.font.name = FONT

    doc.add_paragraph(
        "Sensitivity analysis establishes that, in contrast with SOSAC, TSAC rapidly becomes "
        "a dominant overlay on the IUSAF equity base and requires careful consideration of "
        "balance points. Progressively increasing the TSAC value produces a structural "
        "breakpoint where the IUSAF band ordering is subverted: China (Band 6, lowest IUSAF "
        "weight) receives a higher per-Party allocation than Brazil, India and Mexico (Band 5). "
        "Under pure IUSAF, the allocation ordering is strictly monotonic: "
        "Band 1 > Band 2 > Band 3 > Band 4 > Band 5 > Band 6."
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "The model adopts band-order preservation as its structural constraint: the Gini-minimum "
        "TSAC value is the lowest point where the Gini coefficient is minimised while maintaining "
        "the descending band order (Band 5 mean > Band 6 mean). A Spearman safety floor of ρ = 0.80 "
        "is also maintained but does not bind at any policy-relevant TSAC setting (the Spearman "
        "correlation is ρ = 0.945 at the Gini-minimum, well above the floor)."
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Note: in the discussion below, unless otherwise specified, SOSAC is set at 3% and all "
        "figures refer to the 142 eligible Parties with high-income countries excluded."
    ).paragraph_format.space_after = Pt(12)

    # ---- SECTION 2: Three Balance Points ----
    h2 = doc.add_heading("Three Balance Points", level=2)
    for run in h2.runs:
        run.font.name = FONT

    doc.add_paragraph(
        "Three balance points are presented for potential consideration, each defined by its "
        "relationship to the band-order preservation constraint:"
    ).paragraph_format.space_after = Pt(6)

    # Strict
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("Strict (TSAC = 1.5%, SOSAC = 3%): ")
    run.font.bold = True
    run.font.name = FONT
    run = p.add_run(
        "The IUSAF equity base is dominant for all eligible Parties without exception. "
        "Stewardship supplements never exceed the equity base for any Party. Band order is "
        "preserved with a margin of 19.3% (Band 5 mean $4.92M vs Band 6 mean $3.97M). "
        f"Spearman ρ = {scenarios[0.015]['spearman']:.3f}."
    )
    run.font.name = FONT

    # Gini-minimum
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("Gini-minimum (TSAC = 2.5%, SOSAC = 3%): ")
    run.font.bold = True
    run.font.name = FONT
    run = p.add_run(
        "The point at which the Gini coefficient is lowest while preserving IUSAF band order. "
        "If this balance point is chosen, stewardship may approach but does not exceed the "
        "equity base for developing country Parties with a large landmass. Band order is "
        "preserved with a margin of 5.4% (Band 5 mean $5.44M vs Band 6 mean $5.15M). "
        f"Spearman ρ = {scenarios[0.025]['spearman']:.3f}."
    )
    run.font.name = FONT

    # Band-order boundary
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("Band-order boundary (TSAC ≈ 3.0%, SOSAC = 3%): ")
    run.font.bold = True
    run.font.name = FONT
    run = p.add_run(
        "The structural ceiling at which the descending band order is just overturned: "
        "Band 6 mean allocation ($5.74M) marginally exceeds Band 5 ($5.70M). "
        "Beyond this point, China (Band 6) receives a higher allocation than Brazil, India "
        "and Mexico (Band 5), breaking the IUSAF monotonic requirement. "
        f"Spearman ρ = {scenarios[0.03]['spearman']:.3f}."
    )
    run.font.name = FONT

    doc.add_paragraph()

    # ---- SECTION 3: Stewardship Pool Volumes ----
    h2 = doc.add_heading("Stewardship Pool Volumes", level=2)
    for run in h2.runs:
        run.font.name = FONT

    doc.add_paragraph(
        "In considering these balance points it is important to calculate the volume of funds "
        "that are being allocated for IUSAF and TSAC and SOSAC. Under a US$1 billion fund "
        "93.5–95.5% ($935–955m) will be stably allocated under the IUSAF. The amount under "
        "consideration for allocation under the three balance point scenarios is 4.5–6.5% of "
        "the Cali Fund or $45–65 million. Under these scenarios for 142 developing country "
        "Parties the vast majority would see reallocations of less than $0.5 million. As such, "
        "if IUSAF is accepted as the stable baseline, the question for consideration by Parties "
        "would be how to allocate the $45–65 million stewardship pool. Note that the size of "
        "the stewardship pool will scale in accordance with the overall size of the Cali Fund."
    ).paragraph_format.space_after = Pt(12)

    # ---- SECTION 4: Band-Order Preservation Table ----
    h2 = doc.add_heading("Band-Order Preservation Across TSAC Levels", level=2)
    for run in h2.runs:
        run.font.name = FONT

    doc.add_paragraph(
        "Table A shows how the mean allocation for Band 6 (China) and Band 5 (Brazil, India, "
        "Mexico) changes as TSAC increases, and whether the descending IUSAF band order is "
        "preserved."
    ).paragraph_format.space_after = Pt(6)

    # Table A: Band-order preservation
    headers_a = ['TSAC Level', 'Band 6 mean\n(China)', 'Band 5 mean\n(Brazil, India, Mexico)',
                 'Band 5 vs\nBand 6 margin', 'IUSAF Band\nOrder Preserved?']
    rows_a = []
    for beta in [0, 0.015, 0.025, 0.03, 0.035, 0.092]:
        m = scenarios[beta]
        preserved = (f"YES (margin {m['margin']:.1f}%)" if m['order_ok'] and m['margin'] < 20
                     else "YES" if m['order_ok']
                     else "NO — Band 6 overtakes Band 5")
        rows_a.append([
            m['label'],
            f"USD {m['b6_mean']:.2f}M",
            f"USD {m['b5_mean']:.2f}M",
            f"{m['margin']:+.1f}%",
            preserved,
        ])
    # Row indices for highlighting: 2=Gini-minimum(green), 3=band-order boundary(red)
    styled_table(doc, headers_a, rows_a,
                 col_widths=[Cm(5.5), Cm(3.0), Cm(4.0), Cm(2.5), Cm(4.5)],
                 highlight_rows={2: "green", 3: "red"})

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Table A. Band-order preservation across TSAC levels (SOSAC = 3%)")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.italic = True

    doc.add_paragraph(
        "The key observation in Table A is that at TSAC = 2.5% (the Gini-minimum) the IUSAF "
        "band order is preserved with a margin of 5.4%, but at TSAC = 3.0% Band 6 is receiving "
        "a higher allocation than Band 5, breaking the descending monotonic requirement. "
        "This overturn occurs at a TSAC value 17× lower than the simple model overturn threshold "
        "(TSAC + SOSAC = 50%), underscoring the sensitivity of the band ordering to the TSAC "
        "stewardship overlay."
    ).paragraph_format.space_after = Pt(12)

    # ---- SECTION 5: Break-Point Summary Table ----
    h2 = doc.add_heading("Break-Point Summary", level=2)
    for run in h2.runs:
        run.font.name = FONT

    doc.add_paragraph(
        "Table B summarises the key metrics at each break-point, including the Spearman rank "
        "correlation against the pure IUSAF baseline, the IUSAF share of the fund, and the "
        "structural status of the band ordering."
    ).paragraph_format.space_after = Pt(6)

    headers_b = ['TSAC', 'SOSAC', 'IUSAF %', 'Spearman ρ', 'Band Order', 'What Happens']
    rows_b = [
        ['0%', '3%', '97.0%', '0.977', 'Preserved', 'SOSAC only — modest rank shift among SIDS'],
        ['1.5%', '3%', '95.5%', '0.951', 'Preserved (19.3%)', 'Strict — IUSAF dominant for all Parties'],
        ['2.5%', '3%', '94.5%', '0.945', 'Preserved (5.4%)', 'Gini-minimum — lowest Gini preserving band order'],
        ['3.0%', '3%', '94.0%', '0.929', 'Overturned', 'Band-order boundary — Band 6 > Band 5'],
        ['3.5%', '3%', '93.5%', '0.917', 'Overturned', 'Bounded — band order already overturned'],
        ['9.2%', '3%', '87.8%', '0.696', 'Overturned', 'TSAC component overturn for China'],
    ]
    styled_table(doc, headers_b, rows_b,
                 col_widths=[Cm(1.5), Cm(1.5), Cm(2.0), Cm(2.0), Cm(3.0), Cm(6.0)],
                 highlight_rows={2: "green", 3: "red"})

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Table B. TSAC break-point summary (SOSAC = 3%, Spearman vs pure IUSAF)")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.italic = True

    doc.add_paragraph(
        "It may be noted that the Spearman safety floor of ρ = 0.80 does not bind at any "
        "policy-relevant TSAC setting. The lowest Spearman value in the negotiation range "
        "(TSAC 1.5–3.0%) is ρ = 0.929, well above the floor. The Spearman correlation "
        "declines as TSAC increases through the transition zone but remains high throughout "
        "the policy-relevant range. Band-order preservation, not Spearman, is the binding "
        "structural constraint."
    ).paragraph_format.space_after = Pt(12)

    # ---- SECTION 6: Ranking Trajectory ----
    h2 = doc.add_heading("Ranking Trajectory", level=2)
    for run in h2.runs:
        run.font.name = FONT

    doc.add_paragraph(
        "Table C shows the full ranking trajectory as TSAC increases from 0% to 10% at 0.5 "
        "percentage point intervals, with the Gini coefficient, Spearman correlation, and "
        "band-order status at each step."
    ).paragraph_format.space_after = Pt(6)

    headers_c = ['TSAC %', 'IUSAF %', 'Spearman ρ', 'Gini', 'Band 6\nmean ($M)',
                 'Band 5\nmean ($M)', 'Band Order']
    rows_c = []
    for _, r in trajectory.iterrows():
        rows_c.append([
            f"{r['TSAC %']:.1f}", f"{r['IUSAF %']:.1f}", f"{r['Spearman ρ']:.3f}",
            f"{r['Gini']:.4f}", f"{r['Band 6 mean']:.2f}", f"{r['Band 5 mean']:.2f}",
            str(r['Band order']),
        ])

    # Highlight key rows: TSAC=2.5 (row index 5), TSAC=3.0 (row index 6)
    styled_table(doc, headers_c, rows_c,
                 col_widths=[Cm(1.5), Cm(1.5), Cm(2.0), Cm(1.5), Cm(2.0), Cm(2.0), Cm(2.5)],
                 highlight_rows={5: "green", 6: "red"})

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Table C. Full TSAC ranking trajectory (SOSAC = 3%)")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.italic = True

    doc.add_paragraph(
        "The ranking trajectory shows that the Gini coefficient reaches its minimum at "
        "TSAC = 2.5% (Gini = 0.0886) while band order is still preserved. Beyond the "
        "band-order overturn at TSAC = 3.0%, the Gini continues to decline briefly before "
        "rising again as TSAC dominates. The policy-relevant range lies between the strict "
        "point (TSAC = 1.5%) and the Gini-minimum (TSAC = 2.5%), with the band-order "
        "boundary at TSAC = 3.0% as the structural ceiling."
    ).paragraph_format.space_after = Pt(12)

    # ---- SECTION 7: Beyond the Boundary ----
    h2 = doc.add_heading("Reference Points Beyond the Policy Range", level=2)
    for run in h2.runs:
        run.font.name = FONT

    doc.add_paragraph(
        "For completeness, two further reference points lie beyond the band-order boundary:"
    ).paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("Bounded (TSAC = 3.5%): ")
    run.font.bold = True
    run.font.name = FONT
    run = p.add_run(
        "Under earlier versions of the model this was a named balance point where IUSAF "
        "remained dominant for most Parties but band order was already overturned. Under "
        "Option D this is classified as beyond the structural boundary since Band 6 clearly "
        "exceeds Band 5 (Band 6 mean $6.33M vs Band 5 mean $5.97M). Parties may still wish "
        "to consider this setting if they are willing to accept band-order overturn in exchange "
        "for greater terrestrial stewardship recognition. Spearman ρ = 0.917."
    )
    run.font.name = FONT

    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("TSAC Component Overturn (TSAC ≈ 9.2%): ")
    run.font.bold = True
    run.font.name = FONT
    run = p.add_run(
        "At this point the TSAC component overtakes the IUSAF component for China — meaning "
        "China's allocation is driven more by its land area than by its inverted UN Scale share. "
        "This is far beyond the policy-relevant range. Spearman ρ = 0.696, indicating "
        "substantial rank reordering."
    )
    run.font.name = FONT

    doc.add_paragraph()

    # ---- SECTION 8: Monotonic Order Consideration ----
    h2 = doc.add_heading("Consideration of Band-Order Relaxation", level=2)
    for run in h2.runs:
        run.font.name = FONT

    doc.add_paragraph(
        "A final consideration for Parties is whether the balance points above should preserve "
        "strictly decreasing (monotonic) order in the amounts allocated by IUSAF band or "
        "whether relaxation is permitted in order to adequately recognise the stewardship "
        "responsibilities of some Parties. Under Option D, the Gini-minimum balance point "
        "(TSAC = 2.5%) preserves band order with a margin of 5.4%. If Parties wish to go "
        "beyond this towards greater terrestrial stewardship recognition, they may accept "
        "band-order relaxation, in which case settings up to TSAC = 3.5% (the former "
        "\"Bounded\" point) remain within the range where IUSAF is still the dominant "
        "component for most Parties. However, such settings would mean that Parties in "
        "a higher band (Band 6) could receive allocations higher than Parties in a lower "
        "band (Band 5), breaching the strictly decreasing order of allocations."
    ).paragraph_format.space_after = Pt(12)

    # Save
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    import duckdb
    main()
