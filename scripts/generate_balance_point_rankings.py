"""Generate ranked country tables for the three Option D balance points.

Produces CSV and Word files for:
- Strict (TSAC=1.5%, SOSAC=3%)
- Gini-minimum (TSAC=2.5%, SOSAC=3%)
- Band-order boundary (TSAC=3.0%, SOSAC=3%)

Each table: Rank, Party, Total Allocation, State, IPLC, Income Group, LDC, EU, UN Band
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from cali_model.data_loader import load_data, get_base_data
from cali_model.calculator import calculate_allocations

FONT = "Times New Roman"
HEADER_BG = "D9D9D9"
SIDS_HIGHLIGHT = "E8F5E9"
OUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'model-tables')

FUND = 1_000_000_000
IPLC = 50

SCENARIOS = [
    ('strict', 'Strict (TSAC=1.5%, SOSAC=3%)', 0.015, 0.03),
    ('gini_minimum', 'Gini-minimum (TSAC=2.5%, SOSAC=3%)', 0.025, 0.03),
    ('band_order_boundary', 'Band-order boundary (TSAC=3.0%, SOSAC=3%)', 0.03, 0.03),
]

DISPLAY_COLS = ['party', 'total_allocation', 'state_component', 'iplc_component',
                'WB Income Group', 'is_ldc', 'is_sids', 'is_eu_ms', 'un_band']
HEADERS = ['Rank', 'Party', 'Total\n(USD M)', 'State\n(USD M)', 'IPLC\n(USD M)',
           'Income Group', 'LDC', 'SIDS', 'EU', 'UN Band']


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def generate_scenario(con, name, label, beta, gamma):
    base_df = get_base_data(con)
    df = calculate_allocations(base_df, FUND, IPLC, exclude_high_income=True,
                               tsac_beta=beta, sosac_gamma=gamma, equality_mode=False,
                               un_scale_mode="band_inversion")

    # Filter to eligible only, sort by allocation desc
    eligible = df[df['eligible']].copy()
    eligible = eligible.sort_values(['total_allocation', 'party'], ascending=[False, True]).reset_index(drop=True)
    eligible.index = eligible.index + 1
    eligible.index.name = 'Rank'

    # CSV
    csv_path = os.path.join(OUT_DIR, f'iusaf-{name}-ranked-country.csv')
    csv_df = eligible[DISPLAY_COLS].copy()
    csv_df.insert(0, 'Rank', csv_df.index)
    csv_df.to_csv(csv_path, index=False)
    print(f"Saved: {csv_path}")

    # Word
    doc_path = os.path.join(OUT_DIR, f'iusaf-{name}-ranked-country.docx')
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    title = doc.add_heading(f"IUSAF Allocation: {label}", level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = FONT
        run.font.size = Pt(12)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"142 eligible Parties | Fund size: USD 1,000 M | State/IPLC split: 50/50 | Ranked by total allocation")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.italic = True

    doc.add_paragraph()

    n_rows = len(eligible)
    n_cols = len(HEADERS)
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for j, hdr in enumerate(HEADERS):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for k, line in enumerate(hdr.split("\n")):
            if k > 0:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.name = FONT
            run.font.size = Pt(7)
            run.font.bold = True
        set_cell_shading(cell, HEADER_BG)

    # Data
    for i, (idx, row) in enumerate(eligible.iterrows()):
        is_sids = row.get('is_sids', False)
        values = [
            str(idx),
            str(row['party']),
            f"{row['total_allocation']:,.2f}",
            f"{row['state_component']:,.2f}",
            f"{row['iplc_component']:,.2f}",
            str(row['WB Income Group']),
            'LDC' if row['is_ldc'] else '-',
            'SIDS' if row['is_sids'] else '-',
            'EU' if row['is_eu_ms'] else '-',
            str(row['un_band']).split(':')[0] if pd.notna(row['un_band']) else '',
        ]

        for j, val in enumerate(values):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(7)
            if j == 1 and is_sids:
                run.font.bold = True

        if is_sids:
            for j in range(n_cols):
                set_cell_shading(table.rows[i + 1].cells[j], SIDS_HIGHLIGHT)

    # Column widths
    widths = [Cm(1.0), Cm(4.5), Cm(2.0), Cm(2.0), Cm(2.0), Cm(3.0), Cm(1.5), Cm(2.0), Cm(2.5)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w

    doc.save(doc_path)
    print(f"Saved: {doc_path}")


def main():
    import duckdb
    con = duckdb.connect(database=':memory:')
    load_data(con)
    os.makedirs(OUT_DIR, exist_ok=True)

    for name, label, beta, gamma in SCENARIOS:
        generate_scenario(con, name, label, beta, gamma)


if __name__ == "__main__":
    main()
